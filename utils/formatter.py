from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def format_business_plan(data: dict) -> str:
    return f"""
================ BUSINESS PLAN ================

1. Executive Summary
{data['refined_idea']}

---------------------------------------------

2. Product & Problem
{data['idea_full']}

---------------------------------------------

3. Market Analysis
{data['market_research']}

---------------------------------------------

4. Business Model
{data['business']}

---------------------------------------------

5. Marketing Strategy
{data['growth_strategy']}

---------------------------------------------

=============================================
"""

def add_page_border(doc):
    section = doc.sections[0]
    sectPr = section._sectPr

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')     # line style
        border.set(qn('w:sz'), '24')          # thickness (24 = ~1pt)
        border.set(qn('w:space'), '24')       # distance from text
        border.set(qn('w:color'), '000000')   # black

        pgBorders.append(border)

    sectPr.append(pgBorders)

def save_to_docx(data: dict):
    doc = Document()

    add_page_border(doc) # to add a border to the page

    # ✅ Set margins (narrow)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ✅ Title
    title = doc.add_paragraph()
    title_run = title.add_run("BUSINESS PLAN")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    # 🧩 Helper function
    def add_section(heading, content):
        # Heading
        h = doc.add_paragraph()
        run = h.add_run(heading)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

        # Content
        for line in content.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

    # ✅ Sections
    add_section("1. Executive Summary", data["refined_idea"])
    add_section("2. Product & Problem", data["idea_full"])
    add_section("3. Market Analysis", data["market_research"])
    add_section("4. Business Model", data["business"])
    add_section("5. Marketing Strategy", data["growth_strategy"])

    # # 💾 Save
    # doc.save("business_plan.docx")

    # 🔥 create outputs folder if not exists
    os.makedirs("outputs", exist_ok=True)

    file_path = os.path.join("outputs", "business_plan.docx")
    doc.save(file_path)

    return file_path